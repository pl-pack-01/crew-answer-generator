"""Document ingestion: parse DOCX/PDF into structured FormSchema using Claude API."""

from __future__ import annotations

import json
import os

import anthropic
from docx import Document

from .models import FieldType, FormSchema

EXTRACTION_PROMPT = """You are a document form parser. Analyze the following document content and extract ALL questions/fields into a structured JSON format.

For each question, determine:
- "text": The question or field label
- "field_type": One of: dropdown, multi_select, text, textarea, date, yes_no, number
- "options": For dropdown/multi_select, provide the list of options. If the document implies a limited set (like Yes/No, states, etc.), enumerate them.
- "required": true/false based on whether the field appears mandatory
- "help_text": Any additional instructions or context for the field
- "section": The section/category this question belongs to

Rules for field_type detection:
- If the field has a small fixed set of choices -> dropdown
- If multiple selections are allowed -> multi_select
- If it's a Yes/No or True/False -> yes_no
- If it asks for a date -> date
- If it asks for a number/count/quantity -> number
- If it needs a short answer (name, email, etc.) -> text
- If it needs a long answer (description, notes) -> textarea

Group questions into logical sections based on the document structure.

Return valid JSON matching this structure:
{
  "name": "Form name derived from document",
  "description": "Brief description of the form's purpose",
  "sections": [
    {
      "title": "Section Name",
      "description": "Section description if any",
      "questions": [
        {
          "text": "Question text",
          "field_type": "dropdown",
          "options": ["Option 1", "Option 2"],
          "required": true,
          "help_text": "Additional context",
          "section": "Section Name"
        }
      ]
    }
  ]
}

Return ONLY the JSON, no other text.

Document content:
"""


def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def parse_document_with_claude(document_text: str) -> dict:
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        messages=[
            {"role": "user", "content": EXTRACTION_PROMPT + document_text}
        ],
    )
    response_text = message.content[0].text
    # Strip markdown code fences if present
    if response_text.startswith("```"):
        lines = response_text.split("\n")
        lines = lines[1:]  # remove opening fence
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        response_text = "\n".join(lines)
    return json.loads(response_text)


def ingest_document(file_path: str, source_filename: str) -> FormSchema:
    document_text = extract_text_from_docx(file_path)
    parsed = parse_document_with_claude(document_text)
    schema = FormSchema(
        name=parsed.get("name", source_filename),
        description=parsed.get("description"),
        source_filename=source_filename,
        sections=[],
    )
    for section_data in parsed.get("sections", []):
        from .models import Question, Section
        questions = []
        for q in section_data.get("questions", []):
            ft = q.get("field_type", "text")
            try:
                field_type = FieldType(ft)
            except ValueError:
                field_type = FieldType.TEXT
            questions.append(Question(
                text=q["text"],
                field_type=field_type,
                options=q.get("options", []),
                required=q.get("required", True),
                help_text=q.get("help_text"),
                section=section_data.get("title"),
            ))
        schema.sections.append(Section(
            title=section_data.get("title", "General"),
            description=section_data.get("description"),
            questions=questions,
        ))
    return schema
