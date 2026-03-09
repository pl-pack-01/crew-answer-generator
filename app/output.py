"""Generate filled documents from customer responses."""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.shared import Pt

from .models import FormResponse, FormSchema


def _evaluate_conditions(conditions, answers) -> bool:
    """Check if a question's conditions are met. Returns True if the question should be visible."""
    if not conditions:
        return True
    for cond in conditions:
        current_val = answers.get(cond.question_id)
        if current_val is None:
            return False
        if cond.operator == "equals" and current_val != cond.value:
            return False
        if cond.operator == "not_equals" and current_val == cond.value:
            return False
        if cond.operator == "contains" and cond.value not in (current_val if isinstance(current_val, list) else [current_val]):
            return False
        if cond.operator == "in" and current_val not in cond.value:
            return False
    return True


def generate_filled_docx(schema: FormSchema, response: FormResponse, output_path: str) -> Path:
    """Generate a new DOCX with all questions and their answers filled in."""
    doc = Document()

    # Title
    title = doc.add_heading(schema.name, level=0)
    if schema.description:
        doc.add_paragraph(schema.description)

    if response.customer_name:
        doc.add_paragraph(f"Customer: {response.customer_name}")

    if response.submitted_at:
        doc.add_paragraph(f"Submitted: {response.submitted_at.strftime('%Y-%m-%d %H:%M')}")

    if response.signed_off:
        doc.add_paragraph(f"Signed off: {response.signed_off_at.strftime('%Y-%m-%d %H:%M') if response.signed_off_at else 'Yes'}")

    doc.add_paragraph("")  # spacer

    for section in schema.sections:
        doc.add_heading(section.title, level=1)
        if section.description:
            doc.add_paragraph(section.description)

        for question in section.questions:
            # Determine if this question was visible to the customer
            visible = _evaluate_conditions(question.conditions, response.answers)

            # Question text in bold
            q_para = doc.add_paragraph()
            run = q_para.add_run(question.text)
            run.bold = True
            run.font.size = Pt(11)

            if not visible:
                # Question was hidden by progressive disclosure
                a_para = doc.add_paragraph("N/A — Not applicable based on responses")
                a_para.paragraph_format.left_indent = Pt(18)
                a_para.runs[0].italic = True
            else:
                # Answer
                answer = response.answers.get(question.id, "")
                if isinstance(answer, list):
                    answer = ", ".join(answer)
                a_para = doc.add_paragraph(answer or "(No answer provided)")
                a_para.paragraph_format.left_indent = Pt(18)

    path = Path(output_path)
    doc.save(str(path))
    return path
