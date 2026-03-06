"""Tests for document ingestion."""

import json
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from app.ingestion import extract_text_from_docx, ingest_document, parse_document_with_claude
from app.models import FieldType


@pytest.fixture
def sample_docx(tmp_path):
    """Create a simple DOCX file for testing."""
    path = tmp_path / "test_form.docx"
    doc = Document()
    doc.add_heading("Customer Intake Form", level=0)
    doc.add_paragraph("What is your company name?")
    doc.add_paragraph("Which environment do you use? (Production / Staging / Development)")
    doc.add_paragraph("Do you require SSO? (Yes / No)")
    doc.add_paragraph("Describe your integration needs:")
    doc.save(str(path))
    return str(path)


@pytest.fixture
def sample_docx_with_table(tmp_path):
    """Create a DOCX with a table."""
    path = tmp_path / "table_form.docx"
    doc = Document()
    doc.add_heading("Requirements", level=0)
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Name"
    table.rows[1].cells[1].text = ""
    table.rows[2].cells[0].text = "Region"
    table.rows[2].cells[1].text = ""
    doc.save(str(path))
    return str(path)


MOCK_CLAUDE_RESPONSE = {
    "name": "Customer Intake Form",
    "description": "Customer onboarding questionnaire",
    "sections": [
        {
            "title": "Company Information",
            "description": "Basic company details",
            "questions": [
                {
                    "text": "What is your company name?",
                    "field_type": "text",
                    "options": [],
                    "required": True,
                    "help_text": None,
                },
                {
                    "text": "Which environment do you use?",
                    "field_type": "dropdown",
                    "options": ["Production", "Staging", "Development"],
                    "required": True,
                    "help_text": "Select your primary environment",
                },
            ],
        },
        {
            "title": "Technical Requirements",
            "questions": [
                {
                    "text": "Do you require SSO?",
                    "field_type": "yes_no",
                    "options": [],
                    "required": True,
                    "help_text": None,
                },
                {
                    "text": "Describe your integration needs",
                    "field_type": "textarea",
                    "options": [],
                    "required": False,
                    "help_text": "Provide details about required integrations",
                },
            ],
        },
    ],
}


class TestExtractText:
    def test_extracts_paragraphs(self, sample_docx):
        text = extract_text_from_docx(sample_docx)
        assert "What is your company name?" in text
        assert "Which environment do you use?" in text
        assert "Do you require SSO?" in text

    def test_extracts_tables(self, sample_docx_with_table):
        text = extract_text_from_docx(sample_docx_with_table)
        assert "Name" in text
        assert "Region" in text

    def test_skips_empty_paragraphs(self, tmp_path):
        path = tmp_path / "empty.docx"
        doc = Document()
        doc.add_paragraph("")
        doc.add_paragraph("Real content")
        doc.add_paragraph("")
        doc.save(str(path))
        text = extract_text_from_docx(str(path))
        assert text == "Real content"


class TestParseWithClaude:
    @patch("app.ingestion.anthropic.Anthropic")
    def test_parses_json_response(self, mock_anthropic_cls):
        mock_client = MagicMock()
        mock_anthropic_cls.return_value = mock_client
        mock_message = MagicMock()
        mock_message.content = [MagicMock(text=json.dumps(MOCK_CLAUDE_RESPONSE))]
        mock_client.messages.create.return_value = mock_message

        result = parse_document_with_claude("some document text")
        assert result["name"] == "Customer Intake Form"
        assert len(result["sections"]) == 2

    @patch("app.ingestion.anthropic.Anthropic")
    def test_strips_markdown_fences(self, mock_anthropic_cls):
        mock_client = MagicMock()
        mock_anthropic_cls.return_value = mock_client
        fenced = f"```json\n{json.dumps(MOCK_CLAUDE_RESPONSE)}\n```"
        mock_message = MagicMock()
        mock_message.content = [MagicMock(text=fenced)]
        mock_client.messages.create.return_value = mock_message

        result = parse_document_with_claude("some text")
        assert result["name"] == "Customer Intake Form"


class TestIngestDocument:
    @patch("app.ingestion.parse_document_with_claude")
    def test_creates_schema_from_parsed_data(self, mock_parse, sample_docx):
        mock_parse.return_value = MOCK_CLAUDE_RESPONSE
        schema = ingest_document(sample_docx, "test_form.docx")

        assert schema.name == "Customer Intake Form"
        assert schema.source_filename == "test_form.docx"
        assert len(schema.sections) == 2

        # Check first section
        s1 = schema.sections[0]
        assert s1.title == "Company Information"
        assert len(s1.questions) == 2
        assert s1.questions[0].field_type == FieldType.TEXT
        assert s1.questions[1].field_type == FieldType.DROPDOWN
        assert s1.questions[1].options == ["Production", "Staging", "Development"]

        # Check second section
        s2 = schema.sections[1]
        assert s2.questions[0].field_type == FieldType.YES_NO
        assert s2.questions[1].field_type == FieldType.TEXTAREA
        assert s2.questions[1].required is False

    @patch("app.ingestion.parse_document_with_claude")
    def test_handles_unknown_field_type(self, mock_parse, sample_docx):
        response = {
            "name": "Test",
            "sections": [
                {
                    "title": "General",
                    "questions": [
                        {"text": "Q1", "field_type": "unknown_type", "required": True},
                    ],
                }
            ],
        }
        mock_parse.return_value = response
        schema = ingest_document(sample_docx, "test.docx")
        assert schema.sections[0].questions[0].field_type == FieldType.TEXT
