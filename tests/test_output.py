"""Tests for filled document output generation."""

from datetime import datetime

from docx import Document

from app.models import FieldType, FormResponse, FormSchema, Question, Section
from app.output import generate_filled_docx


def _make_schema():
    return FormSchema(
        id="s1",
        name="Customer Intake",
        description="Intake form for new customers",
        sections=[
            Section(
                id="sec1",
                title="Company Info",
                description="Basic details",
                questions=[
                    Question(id="q1", text="Company Name?", field_type=FieldType.TEXT),
                    Question(id="q2", text="Environment?", field_type=FieldType.DROPDOWN, options=["Prod", "Dev"]),
                    Question(id="q3", text="Features?", field_type=FieldType.MULTI_SELECT, options=["SSO", "MFA"]),
                ],
            ),
            Section(
                id="sec2",
                title="Technical",
                questions=[
                    Question(id="q4", text="Notes?", field_type=FieldType.TEXTAREA),
                ],
            ),
        ],
    )


def _make_response():
    return FormResponse(
        id="r1",
        schema_id="s1",
        schema_version=1,
        customer_name="Acme Corp",
        answers={"q1": "Acme Corp", "q2": "Prod", "q3": ["SSO", "MFA"], "q4": "Need fast onboarding"},
        submitted_at=datetime(2026, 3, 5, 14, 30),
        signed_off=True,
        signed_off_at=datetime(2026, 3, 5, 14, 31),
    )


class TestGenerateFilledDocx:
    def test_creates_valid_docx(self, tmp_path):
        path = tmp_path / "output.docx"
        result = generate_filled_docx(_make_schema(), _make_response(), str(path))
        assert result.exists()
        assert result.suffix == ".docx"

    def test_contains_schema_name(self, tmp_path):
        path = tmp_path / "output.docx"
        generate_filled_docx(_make_schema(), _make_response(), str(path))
        doc = Document(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Customer Intake" in all_text

    def test_contains_customer_name(self, tmp_path):
        path = tmp_path / "output.docx"
        generate_filled_docx(_make_schema(), _make_response(), str(path))
        doc = Document(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Acme Corp" in all_text

    def test_contains_answers(self, tmp_path):
        path = tmp_path / "output.docx"
        generate_filled_docx(_make_schema(), _make_response(), str(path))
        doc = Document(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Prod" in all_text
        assert "SSO, MFA" in all_text
        assert "Need fast onboarding" in all_text

    def test_contains_section_headings(self, tmp_path):
        path = tmp_path / "output.docx"
        generate_filled_docx(_make_schema(), _make_response(), str(path))
        doc = Document(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Company Info" in all_text
        assert "Technical" in all_text

    def test_unanswered_question_shows_placeholder(self, tmp_path):
        path = tmp_path / "output.docx"
        resp = FormResponse(id="r2", schema_id="s1", schema_version=1, answers={})
        generate_filled_docx(_make_schema(), resp, str(path))
        doc = Document(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "(No answer provided)" in all_text

    def test_signed_off_timestamp(self, tmp_path):
        path = tmp_path / "output.docx"
        generate_filled_docx(_make_schema(), _make_response(), str(path))
        doc = Document(str(path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Signed off:" in all_text
        assert "2026-03-05" in all_text
